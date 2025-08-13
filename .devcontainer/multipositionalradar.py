# ----------------------------------------------------------------------
# ⚽ Advanced Multi-Position Player Analysis Tool v4.0 ⚽
#
# This script integrates detailed archetypes for Strikers, Center Midfielders,
# Center Backs, Wingers, and Fullbacks. It allows users to select a position,
# find similar players or upgrades, and generate detailed .docx reports
# with validated metrics and improved visualizations.
# ----------------------------------------------------------------------

# --- 1. IMPORTS ---
import requests
import pandas as pd
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import warnings
import traceback
import matplotlib.pyplot as plt
from rich.console import Console
from rich.progress import Progress
from rich.panel import Panel
from rich.table import Table
from datetime import date

warnings.filterwarnings('ignore')
console = Console()

# --- 2. CORE CONFIGURATION ---
USERNAME = "quadsteriv@gmail.com"
PASSWORD = "SfORY1xR"

LEAGUE_SEASON_MAP = {
    1442: 282, 78: 317, 260: 317, 1581: 317, 179: 317, 129: 317, 1385: 317,
    1778: 282, 1848: 317, 1035: 317, 4: 317, 5: 317, 76: 317, 107: 282,
    51: 317, 1607: 315, 89: 282, 106: 315
}
LEAGUE_NAMES = {
    1442: "Norwegian 1. Division", 78: "Croatian 1. HNL", 260: "Danish 1st Division",
    1581: "Austrian 2. Liga", 179: "German 3. Liga", 129: "French Championnat National",
    1385: "Scottish Championship", 1778: "Irish First Division", 1848: "Polish I Liga",
    1035: "Belgian First Division B", 4: "English League One", 5: "English League Two",
    76: "Belgian First Division A", 107: "LOI Premier Division", 51: "Scottish Premiership",
    1607: "Icelandic Úrvalsdeild", 89: "USL Championship", 106: "Finnish Veikkausliiga"
}

# --- 3. POSITIONAL CONFIGURATIONS ---

# ⚽ STRIKER CONFIG
STRIKER_ARCHETYPES = {
    "Poacher (Fox in the Box)": {
        "description": "Clinical finisher, thrives in the penalty area, instinctive movement, minimal involvement in build-up.",
        "identity_metrics": ['npg_90', 'np_xg_90', 'conversion_ratio', 'touches_inside_box_90', 'shot_touch_ratio', 'np_xg_per_shot'],
        "key_weight": 1.7, "min_percentile_threshold": 60
    },
    "Target Man": {
        "description": "Strong aerial presence, holds up the ball, physical dominance.",
        "identity_metrics": ['aerial_wins_90', 'aerial_ratio', 'fouls_won_90', 'op_xgbuildup_90', 'touches_inside_box_90', 'passes_into_box_90'],
        "key_weight": 1.6, "min_percentile_threshold": 55
    },
    "Complete Forward": {
        "description": "Well-rounded—good finishing, dribbling, link-up, and movement.",
        "identity_metrics": ['npg_90', 'key_passes_90', 'dribbles_90', 'deep_progressions_90', 'op_xgbuildup_90', 'aerial_wins_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    },
    "False 9": {
        "description": "Drops deep into midfield, playmaker-like vision, technical excellence.",
        "identity_metrics": ['op_xgbuildup_90', 'key_passes_90', 'through_balls_90', 'dribbles_90', 'carries_90', 'xa_90'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    },
    "Advanced Forward": {
        "description": "Prioritizes runs in behind, thrives on through balls, pace-driven.",
        "identity_metrics": ['deep_progressions_90', 'through_balls_90', 'np_shots_90', 'touches_inside_box_90', 'npg_90', 'np_xg_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    },
    "Pressing Forward": {
        "description": "Defensive work rate, triggers press, harasses defenders.",
        "identity_metrics": ['pressures_90', 'pressure_regains_90', 'counterpressures_90', 'aggressive_actions_90', 'padj_tackles_90', 'fouls_90'],
        "key_weight": 1.5, "min_percentile_threshold": 55
    },
    "Second Striker (Support Striker)": {
        "description": "Operates just behind main striker, creative link, dribbler.",
        "identity_metrics": ['dribbles_90', 'key_passes_90', 'xa_90', 'touches_inside_box_90', 'npg_90', 'carries_90'],
        "key_weight": 1.5, "min_percentile_threshold": 45
    },
    "Deep-Lying Forward": {
        "description": "Drops into midfield to orchestrate play, but still a striker.",
        "identity_metrics": ['op_xgbuildup_90', 'key_passes_90', 'long_balls_90', 'through_balls_90', 'carries_90', 'passing_ratio'],
        "key_weight": 1.5, "min_percentile_threshold": 45
    },
    "Wide Forward": {
        "description": "Starts wide but cuts inside; often part of a front two or fluid front three.",
        "identity_metrics": ['dribbles_90', 'crosses_90', 'deep_progressions_90', 'touches_inside_box_90', 'npg_90', 'np_shots_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    }
}
STRIKER_RADAR_METRICS = {
    'finishing': { 'name': 'Finishing & Shot Quality', 'color': '#D32F2F', 'metrics': {'npg_90': 'Non-Penalty Goals', 'np_xg_90': 'Non-Penalty xG', 'np_shots_90': 'Shots p90', 'touches_inside_box_90': 'Touches in Box p90', 'conversion_ratio': 'Shot Conversion %', 'np_xg_per_shot': 'Avg. Shot Quality'} },
    'creation': { 'name': 'Creation & Link-Up Play', 'color': '#FF6B35', 'metrics': {'key_passes_90': 'Key Passes p90', 'xa_90': 'xA p90', 'op_passes_into_box_90': 'Passes into Box p90', 'through_balls_90': 'Through Balls p90', 'op_xgbuildup_90': 'xG Buildup p90', 'fouls_won_90': 'Fouls Won p90'} },
    'physicality_pressing': { 'name': 'Physicality & Pressing', 'color': '#4CAF50', 'metrics': {'aerial_wins_90': 'Aerial Duels Won', 'aerial_ratio': 'Aerial Win %', 'pressures_90': 'Pressures p90', 'pressure_regains_90': 'Pressure Regains', 'aggressive_actions_90': 'Aggressive Actions', 'turnovers_90': 'Ball Security (Inv)'} }
}

# ⚽ WINGER CONFIG
WINGER_ARCHETYPES = {
    "Goal-Scoring Winger": {
        "description": "A winger focused on cutting inside to shoot and score goals.",
        "identity_metrics": ['npg_90', 'np_xg_90', 'np_shots_90', 'touches_inside_box_90', 'np_xg_per_shot', 'dribbles_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    },
    "Creative Playmaker": {
        "description": "A winger who creates chances for others through key passes and assists.",
        "identity_metrics": ['xa_90', 'key_passes_90', 'op_passes_into_box_90', 'through_balls_90', 'op_xgbuildup_90', 'deep_progressions_90'],
        "key_weight": 1.5, "min_percentile_threshold": 45
    },
    "Traditional Winger": {
        "description": "Focuses on providing width, dribbling down the line, and delivering crosses.",
        "identity_metrics": ['crosses_90', 'crossing_ratio', 'dribbles_90', 'carry_length', 'deep_progressions_90', 'fouls_won_90'],
        "key_weight": 1.5, "min_percentile_threshold": 40
    }
}
WINGER_RADAR_METRICS = {
    'goal_threat': { 'name': 'Goal Threat', 'color': '#D32F2F', 'metrics': {'npg_90': 'Non-Penalty Goals', 'np_xg_90': 'Non-Penalty xG', 'np_shots_90': 'Shots p90', 'touches_inside_box_90': 'Touches in Box p90', 'conversion_ratio': 'Shot Conversion %', 'np_xg_per_shot': 'Avg. Shot Quality'} },
    'creation_passing': { 'name': 'Creation & Passing', 'color': '#FF6B35', 'metrics': {'key_passes_90': 'Key Passes p90', 'xa_90': 'xA p90', 'op_passes_into_box_90': 'Passes into Box p90', 'through_balls_90': 'Through Balls p90', 'op_xgbuildup_90': 'xG Buildup p90', 'passing_ratio': 'Pass Completion %'} },
    'dribbling_progression': { 'name': 'Dribbling & Progression', 'color': '#9C27B0', 'metrics': {'dribbles_90': 'Successful Dribbles p90', 'dribble_ratio': 'Dribble Success %', 'carries_90': 'Ball Carries p90', 'carry_length': 'Avg. Carry Length', 'deep_progressions_90': 'Deep Progressions p90', 'fouls_won_90': 'Fouls Won p90'} }
}

# ⚽ CENTER MIDFIELDER CONFIG
CM_ARCHETYPES = {
    "Deep-Lying Playmaker (Regista)": {
        "description": "Dictates tempo from deep, excels in progressive passing.",
        "identity_metrics": ['op_xgbuildup_90', 'long_balls_90', 'long_ball_ratio', 'forward_pass_proportion', 'passing_ratio', 'through_balls_90'],
        "key_weight": 1.6, "min_percentile_threshold": 55
    },
    "Box-to-Box Midfielder (B2B)": {
        "description": "Covers large vertical space, contributes in both boxes.",
        "identity_metrics": ['deep_progressions_90', 'carries_90', 'padj_tackles_and_interceptions_90', 'pressures_90', 'npg_90', 'touches_inside_box_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    },
    "Ball-Winning Midfielder (Destroyer)": {
        "description": "Breaks up play, screens defense.",
        "identity_metrics": ['padj_tackles_90', 'padj_interceptions_90', 'pressure_regains_90', 'challenge_ratio', 'aggressive_actions_90', 'fouls_90'],
        "key_weight": 1.6, "min_percentile_threshold": 55
    },
    "Advanced Playmaker (Mezzala)": {
        "description": "Operates in half-spaces, creates in advanced zones.",
        "identity_metrics": ['xa_90', 'key_passes_90', 'op_passes_into_box_90', 'through_balls_90', 'dribbles_90', 'np_shots_90'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    },
    "Transition Midfielder (Tempo Carrier)": {
        "description": "Drives forward in transition, breaks lines with carries.",
        "identity_metrics": ['carries_90', 'carry_length', 'dribbles_90', 'dribble_ratio', 'deep_progressions_90', 'fouls_won_90'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    },
    "Holding Midfielder (Anchor)": {
        "description": "Protects the backline, distributes safely.",
        "identity_metrics": ['padj_interceptions_90', 'passing_ratio', 'op_xgbuildup_90', 'pressures_90', 'challenge_ratio', 'turnovers_90'],
        "key_weight": 1.5, "min_percentile_threshold": 55
    },
    "Attacking Midfielder (8.5 Role)": {
        "description": "Focused on final-third involvement.",
        "identity_metrics": ['npg_90', 'np_xg_90', 'xa_90', 'key_passes_90', 'touches_inside_box_90', 'np_shots_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    }
}
CM_RADAR_METRICS = {
    'ball_winning': { 'name': 'Ball Winning & Defending', 'color': '#D32F2F', 'metrics': {'padj_tackles_90': 'PAdj Tackles', 'padj_interceptions_90': 'PAdj Interceptions', 'pressure_regains_90': 'Pressure Regains', 'challenge_ratio': 'Challenge Success %', 'dribbled_past_90': 'Dribbled Past p90', 'aggressive_actions_90': 'Aggressive Actions'} },
    'progression': { 'name': 'Ball Progression', 'color': '#4CAF50', 'metrics': {'carries_90': 'Ball Carries p90', 'carry_length': 'Avg. Carry Length', 'dribbles_90': 'Successful Dribbles', 'deep_progressions_90': 'Deep Progressions', 'fouls_won_90': 'Fouls Won p90', 'turnovers_90': 'Ball Security (Inv)'} },
    'distribution': { 'name': 'Distribution & Creation', 'color': '#0066CC', 'metrics': {'passing_ratio': 'Pass Completion %', 'forward_pass_proportion': 'Forward Pass %', 'op_xgbuildup_90': 'xG Buildup p90', 'key_passes_90': 'Key Passes p90', 'xa_90': 'xA p90', 'long_balls_90': 'Long Balls p90'} }
}

# ⚽ FULLBACK CONFIG
FULLBACK_ARCHETYPES = {
    "Attacking Fullback": {
        "description": "High attacking output with crosses, key passes, and forward runs into the final third.",
        "identity_metrics": ['xa_90', 'crosses_90', 'op_passes_into_box_90', 'deep_progressions_90', 'key_passes_90', 'op_xgbuildup_90'],
        "key_weight": 1.5, "min_percentile_threshold": 40
    },
    "Defensive Fullback": {
        "description": "Solid defensive foundation with tackles, interceptions, and aerial duels.",
        "identity_metrics": ['padj_tackles_and_interceptions_90', 'challenge_ratio', 'aggressive_actions_90', 'pressures_90', 'aerial_wins_90', 'aerial_ratio'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    },
     "Modern Wingback": {
        "description": "High energy player who covers huge distances, contributing in all phases of play.",
        "identity_metrics": ['deep_progressions_90', 'crosses_90', 'dribbles_90', 'padj_tackles_and_interceptions_90', 'pressures_90', 'xa_90'],
        "key_weight": 1.6, "min_percentile_threshold": 50
    }
}
FULLBACK_RADAR_METRICS = {
    'defensive_actions': { 'name': 'Defensive Actions', 'color': '#00BCD4', 'metrics': {'padj_tackles_and_interceptions_90': 'P.Adj Tackles+Ints p90', 'challenge_ratio': 'Defensive Duel Win %', 'aggressive_actions_90': 'Aggressive Actions p90', 'aerial_wins_90': 'Aerial Duels Won p90', 'aerial_ratio': 'Aerial Win %', 'dribbled_past_90': 'Times Dribbled Past p90'} },
    'progression_creation': { 'name': 'Progression & Creation', 'color': '#FF6B35', 'metrics': {'deep_progressions_90': 'Deep Progressions p90', 'carry_length': 'Avg. Carry Length', 'dribbles_90': 'Successful Dribbles p90', 'dribble_ratio': 'Dribble Success %', 'op_passes_into_box_90': 'Open Play Passes into Box', 'xa_90': 'xA p90'} },
    'work_rate_security': { 'name': 'Work Rate & Security', 'color': '#4CAF50', 'metrics': {'pressures_90': 'Pressures p90', 'pressure_regains_90': 'Pressure Regains p90', 'counterpressures_90': 'Counterpressures p90', 'fouls_won_90': 'Fouls Won p90', 'turnovers_90': 'Ball Security (Inv)', 'op_xgbuildup_90': 'xG Buildup p90'} }
}

# ⚽ CENTER BACK CONFIG
CB_ARCHETYPES = {
    "Ball-Playing Defender": {
        "description": "Comfortable in possession, initiates attacks from the back with progressive passing.",
        "identity_metrics": ['op_xgbuildup_90', 'passing_ratio', 'long_balls_90', 'long_ball_ratio', 'forward_pass_proportion', 'carries_90'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    },
    "Stopper": {
        "description": "Aggressive defender who steps out to challenge attackers and win the ball high up the pitch.",
        "identity_metrics": ['aggressive_actions_90', 'padj_tackles_90', 'challenge_ratio', 'pressures_90', 'aerial_wins_90', 'fouls_90'],
        "key_weight": 1.6, "min_percentile_threshold": 55
    },
    "Covering Defender": {
        "description": "Reads the game well, relying on positioning and interceptions to sweep up behind the defensive line.",
        "identity_metrics": ['padj_interceptions_90', 'padj_clearances_90', 'dribbled_past_90', 'pressure_regains_90', 'aerial_ratio', 'passing_ratio'],
        "key_weight": 1.5, "min_percentile_threshold": 50
    }
}
CB_RADAR_METRICS = {
    'ground_defending': { 'name': 'Ground Duels & Defending', 'color': '#D32F2F', 'metrics': {'padj_tackles_90': 'PAdj Tackles', 'padj_interceptions_90': 'PAdj Interceptions', 'aggressive_actions_90': 'Aggressive Actions', 'challenge_ratio': 'Challenge Success %', 'pressures_90': 'Pressures p90', 'dribbled_past_90': 'Dribbled Past p90'} },
    'aerial_duels': { 'name': 'Aerial Duels', 'color': '#4CAF50', 'metrics': {'aerial_wins_90': 'Aerial Duels Won', 'aerial_ratio': 'Aerial Win %', 'padj_clearances_90': 'PAdj Clearances', 'fouls_90': 'Fouls Committed', 'challenge_ratio': 'Challenge Success %', 'aggressive_actions_90': 'Aggressive Actions'} },
    'passing_progression': { 'name': 'Passing & Progression', 'color': '#0066CC', 'metrics': {'passing_ratio': 'Pass Completion %', 'pass_length': 'Avg. Pass Length', 'long_balls_90': 'Long Balls p90', 'long_ball_ratio': 'Long Ball Accuracy %', 'forward_pass_proportion': 'Forward Pass %', 'op_xgbuildup_90': 'xG Buildup p90'} }
}

# Master configuration dictionary
POSITIONAL_CONFIGS = {
    "Fullback": {"archetypes": FULLBACK_ARCHETYPES, "radars": FULLBACK_RADAR_METRICS, "positions": ['Right Back', 'Left Back', 'Right Wing Back', 'Left Wing Back']},
    "Center Back": {"archetypes": CB_ARCHETYPES, "radars": CB_RADAR_METRICS, "positions": ['Center Back', 'Left Centre Back', 'Right Centre Back']},
    "Center Midfielder": {"archetypes": CM_ARCHETYPES, "radars": CM_RADAR_METRICS, "positions": ['Defensive Midfield', 'Center Defensive Midfield', 'Center Midfield', 'Right Centre Midfielder', 'Left Centre Midfielder']},
    "Winger": {"archetypes": WINGER_ARCHETYPES, "radars": WINGER_RADAR_METRICS, "positions": ['Right Wing', 'Left Wing', 'Right Midfield', 'Left Midfield']},
    "Striker": {"archetypes": STRIKER_ARCHETYPES, "radars": STRIKER_RADAR_METRICS, "positions": ['Attacking Midfield', 'Center Forward', 'Secondary Striker']}
}

# Generate a unified list of all metrics needed for percentiling
ALL_METRICS_TO_PERCENTILE = sorted(list(set(
    metric for pos_config in POSITIONAL_CONFIGS.values()
    for archetype in pos_config['archetypes'].values() for metric in archetype['identity_metrics']
) | set(
    metric for pos_config in POSITIONAL_CONFIGS.values()
    for radar in pos_config['radars'].values() for metric in radar['metrics'].keys()
)))


# --- 4. DATA HANDLING & ANALYSIS FUNCTIONS ---

def sanitize_filename(name):
    """Removes invalid characters from a string to make it a valid filename."""
    name = name.replace(' ', '_')
    return "".join([c for c in name if c.isalpha() or c.isdigit() or c in ['_','.']]).rstrip()

def get_all_leagues_data(auth_credentials):
    """Downloads player statistics from all leagues defined in LEAGUE_SEASON_MAP."""
    all_dfs = []
    console.rule("[bold blue]Loading Player Data[/bold blue]")
    with Progress() as progress:
        task = progress.add_task("[green]Downloading...", total=len(LEAGUE_SEASON_MAP))
        for league_id, season_id in LEAGUE_SEASON_MAP.items():
            league_name = LEAGUE_NAMES.get(league_id, f"League {league_id}")
            progress.update(task, description=f"[green]Downloading [cyan]{league_name}[/cyan]...")
            try:
                url = f"https://data.statsbombservices.com/api/v1/competitions/{league_id}/seasons/{season_id}/player-stats"
                response = requests.get(url, auth=auth_credentials)
                response.raise_for_status()
                df_league = pd.json_normalize(response.json())
                df_league['league_name'] = league_name
                all_dfs.append(df_league)
            except Exception as e:
                console.log(f"[red]Failed to load {league_name}: {str(e)[:70]}...[/red]")
            progress.update(task, advance=1)
    if not all_dfs: raise Exception("No league data could be loaded.")
    return pd.concat(all_dfs, ignore_index=True)

def calculate_age_from_birth_date(birth_date_str):
    """Calculates player age from a birth date string."""
    if pd.isna(birth_date_str): return None
    try:
        birth_date = pd.to_datetime(birth_date_str).date()
        today = date.today()
        return today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
    except (ValueError, TypeError): return None

def process_and_percentile_data(df):
    """Processes raw data, calculates ages, and computes position-specific percentiles."""
    console.rule("[bold blue]Processing Player Data[/bold blue]")
    df_processed = df.copy()
    df_processed.columns = [c.replace('player_season_', '') for c in df_processed.columns]
    df_processed['age'] = df_processed['birth_date'].apply(calculate_age_from_birth_date)
    
    # Calculate a combined tackles and interceptions metric if not present
    if 'padj_tackles_90' in df_processed.columns and 'padj_interceptions_90' in df_processed.columns:
        df_processed['padj_tackles_and_interceptions_90'] = df_processed['padj_tackles_90'] + df_processed['padj_interceptions_90']
    
    # Calculate percentiles within each major positional group
    for pos_group, config in POSITIONAL_CONFIGS.items():
        console.log(f"-> Calculating percentiles for [bold cyan]{pos_group}s[/bold cyan]...")
        pos_mask = df_processed['primary_position'].isin(config['positions'])
        
        for metric in ALL_METRICS_TO_PERCENTILE:
            if metric in df_processed.columns:
                metric_data = df_processed.loc[pos_mask, metric]
                if pd.api.types.is_numeric_dtype(metric_data) and not metric_data.empty:
                    pct_col = f'{metric}_pct'
                    negative_stats = ['turnovers_90', 'dispossessions_90', 'dribbled_past_90', 'fouls_90']
                    ranks = metric_data.rank(pct=True) * 100
                    df_processed.loc[pos_mask, pct_col] = 100 - ranks if metric in negative_stats else ranks

    console.log("✅ Data processing and percentiling complete.")
    return df_processed

def find_player_by_name(df, player_name):
    """Finds a player by exact or partial name match."""
    exact_matches = df[df['player_name'].str.lower() == player_name.lower()]
    if not exact_matches.empty: return exact_matches.iloc[0].copy()
    
    partial_matches = df[df['player_name'].str.lower().str.contains(player_name.lower(), na=False)]
    if not partial_matches.empty:
        console.print(f"\n[yellow]Player '{player_name}' not found. Did you mean one of these?[/yellow]")
        for _, row in partial_matches[['player_name', 'team_name']].head(5).iterrows():
            console.print(f"  - {row['player_name']} ({row['team_name']})")
    else:
        console.print(f"\n[red]Player '{player_name}' not found.[/red]")
    return None

def detect_player_archetype(target_player, archetypes):
    """Determines the most likely archetype for a player based on their stats."""
    archetype_scores = {}
    for name, config in archetypes.items():
        metrics = [f"{m}_pct" for m in config['identity_metrics']]
        valid_metrics = [m for m in metrics if m in target_player.index and pd.notna(target_player[m])]
        score = target_player[valid_metrics].mean() if valid_metrics else 0
        archetype_scores[name] = score
    
    return max(archetype_scores, key=archetype_scores.get) if archetype_scores else None, archetype_scores

def find_matches(target_player, pool_df, archetype_config, search_mode='similar', min_minutes=500):
    """Finds similar players or upgrades based on the selected archetype and search mode."""
    console.rule(f"[bold blue]Searching for Potential {search_mode.title()}s[/bold blue]")
    
    key_identity_metrics = archetype_config['identity_metrics']
    key_weight = archetype_config['key_weight']
    min_percentile = archetype_config['min_percentile_threshold']
    
    percentile_metrics = [f'{m}_pct' for m in key_identity_metrics]
    
    pool_df = pool_df[(pool_df['minutes'] >= min_minutes) & (pool_df['player_id'] != target_player['player_id'])].dropna(subset=percentile_metrics).copy()
    if pool_df.empty:
        console.log("[yellow]No suitable players found in the pool.[/yellow]")
        return pd.DataFrame()

    console.log(f"Comparing against a pool of [bold]{len(pool_df)}[/] qualified players...")

    target_vector = target_player[percentile_metrics].values.reshape(1, -1)
    pool_matrix = pool_df[percentile_metrics].values
    
    weights = np.full(len(key_identity_metrics), key_weight)
    target_vector_w = target_vector * weights
    pool_matrix_w = pool_matrix * weights
    
    similarities = cosine_similarity(target_vector_w, pool_matrix_w)
    pool_df['similarity_score'] = similarities[0] * 100
    
    initial_count = len(pool_df)
    for metric in key_identity_metrics:
        pool_df = pool_df[pool_df[f"{metric}_pct"] >= min_percentile]
    
    console.log(f"Initial list of {initial_count} players found.")
    console.log(f"-> Filtered to {len(pool_df)} players meeting the minimum percentile in key areas.")
    
    if search_mode == 'upgrade':
        pool_df['upgrade_score'] = pool_df[percentile_metrics].mean(axis=1)
        final_results = pool_df.sort_values('upgrade_score', ascending=False)
    else:
        final_results = pool_df.sort_values('similarity_score', ascending=False)
        
    console.log("✅ Analysis complete.")
    return final_results


# --- 5. REPORTING & VISUALIZATION FUNCTIONS ---

def set_cell_style(cell, text, bold=False, font_size=10, align='CENTER'):
    """Helper function to style cells in a .docx table."""
    p = cell.paragraphs[0]
    p.text = str(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'CENTER' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.font.bold = bold

def create_enhanced_radar_chart(player_data, reference_player, radar_config, filepath):
    """Creates an improved radar chart and saves it to a file."""
    plt.style.use('seaborn-v0_8-notebook')
    metrics_dict = radar_config['metrics']
    labels = ['\n'.join(l.split()) for l in metrics_dict.values()]
    num_vars = len(labels)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist() + [0]
    
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('#F5F5F5')
    
    def get_percentiles(player, metrics):
        values = [max(0, min(100, player.get(f'{m}_pct', 50))) for m in metrics.keys()]
        return values + [values[0]]
        
    player_values = get_percentiles(player_data, metrics_dict)
    player_avg = np.mean(player_values[:-1])
    player_legend = f"{player_data['player_name']} (Avg: {player_avg:.0f}th %ile)"
    
    ax.set_rgrids([20, 40, 60, 80], angle=180)
    ax.set_ylim(0, 105)
    ax.grid(True, color='grey', linestyle='--', linewidth=0.5, alpha=0.5)
    
    ax.fill(angles, player_values, color=radar_config['color'], alpha=0.3, zorder=5)
    ax.plot(angles, player_values, color=radar_config['color'], linewidth=2.5, zorder=6, label=player_legend)
    
    for i, value in enumerate(player_values[:-1]):
        angle = angles[i]
        ax.text(angle, value + 7, f"{value:.0f}", ha='center', va='center', fontweight='bold', size=9,
                color='black', bbox=dict(facecolor='white', alpha=0.7, edgecolor='none', boxstyle='round,pad=0.2'))
    
    if reference_player is not None:
        ref_values = get_percentiles(reference_player, metrics_dict)
        ref_avg = np.mean(ref_values[:-1])
        ref_legend = f"Target: {reference_player['player_name']} (Avg: {ref_avg:.0f}th %ile)"
        ax.plot(angles, ref_values, color='#4A90E2', linewidth=2, zorder=4, linestyle='--', label=ref_legend)
    
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, size=10, fontweight='bold')
    
    ax.set_title(f"{radar_config['name']} | {player_data['player_name']}", size=16, fontweight='bold', y=1.12)
    ax.legend(loc='upper right', bbox_to_anchor=(1.5, 1.15))
    
    plt.savefig(filepath, bbox_inches='tight', dpi=300)
    plt.close()

def create_report_document(target_player, top_matches, other_matches, archetype_dna, search_config, target_radars, comp_radars):
    """Assembles the final .docx report."""
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.add_heading(f'{search_config["position"]} {search_config["mode"].title()} Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(f'Target Player: {target_player["player_name"]}', level=1)

    age = target_player.get('age')
    age_str = f"{int(age)}" if pd.notna(age) else "N/A"
    doc.add_paragraph(f"**Age:** {age_str} | **Team:** {target_player['team_name']} | **League:** {target_player['league_name']}")

    doc.add_heading('Archetype DNA & Search Filters', level=2)
    dna_table = doc.add_table(rows=1, cols=2)
    dna_table.style = 'Table Grid'
    set_cell_style(dna_table.cell(0, 0), "Archetype", bold=True)
    set_cell_style(dna_table.cell(0, 1), "Affinity Score", bold=True)
    for name, score in sorted(archetype_dna.items(), key=lambda item: item[1], reverse=True):
        cells = dna_table.add_row().cells
        set_cell_style(cells[0], name)
        set_cell_style(cells[1], f"{score:.1f}")

    doc.add_paragraph(f"\n**Detected Archetype:** {search_config['archetype']}\n**Search Mode:** {search_config['mode'].title()}")

    doc.add_heading('Target Player Performance Radars', level=2)
    num_radars = len(target_radars)
    rows = (num_radars + 2) // 3
    radar_table = doc.add_table(rows=rows, cols=3)
    for i, path in enumerate(target_radars):
        p = radar_table.cell(i // 3, i % 3).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(2.3))

    if not top_matches.empty:
        doc.add_page_break()
        doc.add_heading(f"Top {len(top_matches)} Matches", level=1)
        for i, (_, player) in enumerate(top_matches.iterrows()):
            age = player.get('age')
            age_str = f"{int(age)}" if pd.notna(age) else "N/A"
            score_col = 'upgrade_score' if search_config["mode"] == 'upgrade' else 'similarity_score'
            score_label = "Upgrade Score" if search_config["mode"] == 'upgrade' else "Similarity"
            doc.add_heading(f"#{i+1}: {player['player_name']} ({score_label}: {player[score_col]:.1f})", level=2)
            doc.add_paragraph(f"**Age:** {age_str} | **Team:** {player['team_name']} | **League:** {player['league_name']}")
            
            comp_radars_table = doc.add_table(rows=rows, cols=3)
            player_radar_paths = comp_radars[i]
            for j, path in enumerate(player_radar_paths):
                p = comp_radars_table.cell(j // 3, j % 3).paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(2.3))
            if i < len(top_matches) - 1: doc.add_page_break()

    if not other_matches.empty:
        doc.add_page_break()
        doc.add_heading("Other Notable Matches", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['Rank', 'Player', 'Age', 'Team', 'League']
        for i, h in enumerate(headers): set_cell_style(table.cell(0, i), h, bold=True)
        for i, (_, row) in enumerate(other_matches.iterrows()):
            cells = table.add_row().cells
            set_cell_style(cells[0], str(i + len(top_matches) + 1))
            set_cell_style(cells[1], row['player_name'], align='LEFT')
            age = row.get('age')
            set_cell_style(cells[2], f"{int(age)}" if pd.notna(age) else "N/A")
            set_cell_style(cells[3], row['team_name'], align='LEFT')
            set_cell_style(cells[4], row['league_name'], align='LEFT')
    
    return doc


# --- 6. MAIN EXECUTION SCRIPT ---

def main():
    """Main function to run the player analysis tool."""
    console.rule("[bold magenta]⚽ Advanced Multi-Position Player Analysis Tool v4.0 ⚽[/bold magenta]")
    try:
        raw_data = get_all_leagues_data((USERNAME, PASSWORD))
        processed_data = process_and_percentile_data(raw_data)
        
        while True:
            console.print("\n" + "="*80)
            
            console.print("[bold]Select a position to analyze:[/bold]")
            pos_options = list(POSITIONAL_CONFIGS.keys())
            for i, pos in enumerate(pos_options):
                console.print(f"  ({i+1}) {pos}")
            pos_choice = console.input("Enter choice [1]: ").strip() or "1"
            
            try:
                selected_position_name = pos_options[int(pos_choice) - 1]
                config = POSITIONAL_CONFIGS[selected_position_name]
                archetypes = config["archetypes"]
                radar_metrics = config["radars"]
                position_pool = processed_data[processed_data['primary_position'].isin(config['positions'])]
            except (ValueError, IndexError):
                console.print("[red]Invalid selection. Please try again.[/red]")
                continue

            player_name = console.input(f"Enter target {selected_position_name}'s full name (or 'quit' to exit): ").strip()
            if player_name.lower() == 'quit': break
            if not player_name: continue

            target_player = find_player_by_name(processed_data, player_name)
            if target_player is None: continue
            
            if target_player['primary_position'] not in config['positions']:
                 console.print(f"[yellow]Warning: {target_player['player_name']}'s primary position is {target_player['primary_position']}, but you are searching in the {selected_position_name} category.[/yellow]")


            detected_archetype, archetype_scores = detect_player_archetype(target_player, archetypes)
            if not detected_archetype:
                console.print(f"[yellow]Could not reliably determine an archetype for {target_player['player_name']}.[/yellow]")
                continue
            
            dna_table = Table(title="Archetype DNA", show_header=True, header_style="bold magenta")
            dna_table.add_column("Archetype", style="cyan")
            dna_table.add_column("Affinity Score", style="green")
            for name, score in sorted(archetype_scores.items(), key=lambda item: item[1], reverse=True):
                dna_table.add_row(name, f"{score:.1f}")
            console.print(dna_table)
            
            archetype_config = archetypes[detected_archetype]

            search_mode = console.input("Select search mode: (1) Find Similar or (2) Find Upgrade [1]: ").strip() or "1"
            search_mode = 'upgrade' if search_mode == '2' else 'similar'
            
            matches = find_matches(target_player, position_pool, archetype_config, search_mode)
            if matches.empty:
                console.print("[bold yellow]No players found who meet the specified criteria.[/bold yellow]")
                continue
            
            top_5_matches = matches.head(5)
            other_matches = matches.iloc[5:20]
            
            console.print(Panel(f"Top 5 Matches for [bold cyan]{target_player['player_name']}[/bold cyan]", title="[bold green]Results[/bold green]"))
            display_cols = ['player_name', 'age', 'similarity_score', 'team_name', 'league_name']
            if search_mode == 'upgrade':
                display_cols.insert(2, 'upgrade_score')
            console.print(top_5_matches[display_cols].rename(columns=lambda c: c.replace('_', ' ').title()).to_string(index=False))
            
            run_report = console.input("\nGenerate a full .docx report? (y/N): ").strip().lower()
            if run_report == 'y':
                console.rule("[bold blue]Generating Full .docx Report[/bold blue]")
                num_charts = len(radar_metrics) * (1 + len(top_5_matches))
                console.print(f"[bold yellow]This will generate {num_charts} radar charts and may take a minute...[/bold yellow]")
                
                export_folder = "player_reports"
                os.makedirs(export_folder, exist_ok=True)
                safe_target_name = sanitize_filename(target_player['player_name'])
                
                target_radars, comp_radars = [], []
                with Progress() as progress:
                    task = progress.add_task("[magenta]Generating radars...", total=num_charts)
                    
                    for _, r_config in radar_metrics.items():
                        path = os.path.join(export_folder, sanitize_filename(f"radar_solo_{safe_target_name}_{r_config['name']}.png"))
                        create_enhanced_radar_chart(target_player, None, r_config, path)
                        target_radars.append(path)
                        progress.update(task, advance=1)
                    
                    for _, player_row in top_5_matches.iterrows():
                        player_radars = []
                        for _, r_config in radar_metrics.items():
                            path = os.path.join(export_folder, sanitize_filename(f"radar_comp_{player_row['player_name']}_{r_config['name']}.png"))
                            create_enhanced_radar_chart(player_row, target_player, r_config, path)
                            player_radars.append(path)
                            progress.update(task, advance=1)
                        comp_radars.append(player_radars)

                search_config = {'mode': search_mode, 'archetype': detected_archetype, 'position': selected_position_name}
                doc = create_report_document(target_player, top_5_matches, other_matches, archetype_scores, search_config, target_radars, comp_radars)
                doc_path = os.path.join(export_folder, f"{search_config['position']}_{search_mode.title()}_Report_{safe_target_name}.docx")
                doc.save(doc_path)
                
                for path in target_radars + [p for sublist in comp_radars for p in sublist]:
                    if os.path.exists(path): os.remove(path)

                console.print(Panel(f"✅ Report for [bold]{target_player['player_name']}[/bold] generated!", title="[bold green]Report Complete[/bold green]"))
                console.print(f"   [link=file://{os.path.abspath(doc_path)}]{os.path.abspath(doc_path)}[/link]")

    except KeyboardInterrupt:
        console.print("\n[yellow]Analysis interrupted by user.[/yellow]")
    except Exception:
        console.print(f"\n[bold red]An unexpected error occurred:[/bold red]")
        traceback.print_exc()
    finally:
        console.rule("[bold magenta]Session Complete[/bold magenta]")

if __name__ == "__main__":
    main()
    
    
    
