import streamlit as st
import requests
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from collections import defaultdict
from docx import Document
from io import BytesIO

# --- 1. APP CONFIGURATION ---

st.set_page_config(
    page_title="St Johnstone Radar Analysis",
    page_icon="⚽",
    layout="wide"
)

# Your StatsBomb Credentials
USERNAME = "quadsteriv@gmail.com"
PASSWORD = "SfORY1xR"


# --- 2. METRICS DEFINITION ---
WINGER_RADAR_METRICS = {
    'attacking_output': {
        'name': 'Attacking Output', 'color': '#D32F2F', 'metrics': {
            'npg_90': 'Non-Penalty Goals',
            'np_xg_90': 'Non-Penalty xG',
            'np_shots_90': 'Shots p90',
            'touches_inside_box_90': 'Touches in Box p90',
            'conversion_ratio': 'Shot Conversion %',
            'np_xg_per_shot': 'Avg. Shot Quality'
        }
    },
    'passing_creation': {
        'name': 'Passing & Creation', 'color': '#FF6B35', 'metrics': {
            'key_passes_90': 'Key Passes p90',
            'xa_90': 'xA p90',
            'op_passes_into_box_90': 'Passes into Box p90',
            'through_balls_90': 'Through Balls p90',
            'op_xgbuildup_90': 'xG Buildup p90',
            'passing_ratio': 'Pass Completion %'
        }
    },
    'dribbling_progression': {
        'name': 'Dribbling & Progression', 'color': '#9C27B0', 'metrics': {
            'dribbles_90': 'Successful Dribbles p90',
            'dribble_ratio': 'Dribble Success %',
            'carries_90': 'Ball Carries p90',
            'carry_length': 'Avg. Carry Length',
            'deep_progressions_90': 'Deep Progressions p90',
            'fouls_won_90': 'Fouls Won p90'
        }
    },
    'crossing': {
        'name': 'Crossing Profile', 'color': '#00BCD4', 'metrics': {
            'crosses_90': 'Completed Crosses p90',
            'crossing_ratio': 'Cross Completion %',
            'box_cross_ratio': '% of Box Passes that are Crosses',
            'sp_passes_into_box_90': 'Set Piece Passes into Box',
            'op_passes_into_box_90': 'Open Play Passes into Box',
            'sp_assists_90': 'Set-Piece Assists p90'
        }
    },
    'defensive_work': {
        'name': 'Pressing & Defensive Work', 'color': '#4CAF50', 'metrics': {
            'pressures_90': 'Pressures p90',
            'pressure_regains_90': 'Pressure Regains p90',
            'counterpressures_90': 'Counterpressures p90',
            'padj_tackles_90': 'P.Adj Tackles p90',
            'padj_interceptions_90': 'P.Adj Interceptions p90',
            'aggressive_actions_90': 'Aggressive Actions p90'
        }
    },
    'physical': {
        'name': 'Physical Profile', 'color': '#607D8B', 'metrics': {
            'aerial_wins_90': 'Aerial Duels Won p90',
            'aerial_ratio': 'Aerial Win %',
            'challenge_ratio': 'Defensive Duel Win %',
            'fouls_90': 'Fouls Committed p90',
            'turnovers_90': 'Turnovers p90 (Ball Security)',
            'dribbled_past_90': 'Times Dribbled Past p90'
        }
    }
}
ALL_METRICS_TO_PERCENTILE = sorted(list(set(
    metric for radar in WINGER_RADAR_METRICS.values() for metric in radar['metrics'].keys()
)))


# --- 3. DATA LOADING AND PROCESSING ---
@st.cache_data
def load_all_data(username, password):
    """Fetches all competition and player data from the API and processes it."""
    comp_url = "https://data.statsbombservices.com/api/v4/competitions"
    try:
        resp = requests.get(comp_url, auth=(username, password))
        resp.raise_for_status()
        api_data = resp.json()
    except requests.exceptions.RequestException as e:
        st.error(f"API Error fetching competitions: {e}")
        return None, None

    leagues = {}
    for item in api_data:
        comp_id = item.get('competition_id')
        if comp_id not in leagues:
            leagues[comp_id] = {
                'name': item.get('competition_name'),
                'seasons': {}
            }
        leagues[comp_id]['seasons'][item.get('season_id')] = item.get('season_name')

    all_dfs = []
    for comp_id, seasons in leagues.items():
        for season_id in seasons['seasons'].keys():
            try:
                player_stats_url = f"https://data.statsbombservices.com/api/v1/competitions/{comp_id}/seasons/{season_id}/player-stats"
                response = requests.get(player_stats_url, auth=(username, password))
                response.raise_for_status()
                df_league = pd.json_normalize(response.json())
                all_dfs.append(df_league)
            except requests.exceptions.RequestException:
                continue
    
    if not all_dfs:
        st.error("Could not load any player data. Check API credentials and access.")
        return None, None

    df = pd.concat(all_dfs, ignore_index=True)
    df.columns = [c.replace('player_season_', '') for c in df.columns]
    for metric in ALL_METRICS_TO_PERCENTILE:
        if metric in df.columns and pd.api.types.is_numeric_dtype(df[metric]):
            pct_col = f'{metric}_pct'
            if metric in ['turnovers_90', 'dribbled_past_90', 'fouls_90']:
                df[pct_col] = 100 - (df[metric].rank(pct=True) * 100)
            else:
                df[pct_col] = df[metric].rank(pct=True) * 100
    
    return leagues, df


# --- 4. VISUALIZATION AND DOCUMENT CREATION ---

def create_word_doc_stream(player1_name, player2_name, chart_name, fig):
    """Creates a Word document with a title and a chart, and returns it as an in-memory stream."""
    image_stream = BytesIO()
    fig.savefig(image_stream, format='png', dpi=300, bbox_inches='tight')
    image_stream.seek(0)

    doc = Document()
    doc.add_heading(f"{chart_name} Comparison", level=1)
    doc.add_heading(f"{player1_name} vs. {player2_name}", level=2)
    doc.add_picture(image_stream)

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    
    return doc_stream

def create_comparison_radar_chart(player1_data, player2_data, radar_config):
    """Generates a single matplotlib radar chart for comparison with data labels."""
    plt.style.use('seaborn-v0_8-darkgrid')
    metrics_dict = radar_config['metrics']
    labels = ['\n'.join(l.split()) for l in metrics_dict.values()]
    num_vars = len(labels)
    
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
    fig.set_facecolor('#121212')
    ax.set_facecolor('#121212')

    def get_percentiles(player, metrics):
        values = [player.get(f'{m}_pct', 0) for m in metrics.keys()]
        values += values[:1]
        return values

    # --- Player 1 Plotting ---
    p1_values = get_percentiles(player1_data, metrics_dict)
    p1_color = '#00f2ff'
    ax.fill(angles, p1_values, color=p1_color, alpha=0.5)
    ax.plot(angles, p1_values, color=p1_color, linewidth=2, label=f"{player1_data['player_name']}")
    
    # Add text labels for Player 1
    for i, (angle, value) in enumerate(zip(angles[:-1], p1_values[:-1])):
        ax.text(angle, value + 5, f"{int(value)}", color=p1_color, ha='center', va='center', fontsize=10, weight='bold')

    # --- Player 2 Plotting ---
    p2_values = get_percentiles(player2_data, metrics_dict)
    p2_color = '#ff0052'
    ax.fill(angles, p2_values, color=p2_color, alpha=0.5)
    ax.plot(angles, p2_values, color=p2_color, linewidth=2, label=f"{player2_data['player_name']}")

    # Add text labels for Player 2
    for i, (angle, value) in enumerate(zip(angles[:-1], p2_values[:-1])):
        ax.text(angle, value + 5, f"{int(value)}", color=p2_color, ha='center', va='center', fontsize=10, weight='bold')

    # --- Chart Styling ---
    ax.set_ylim(0, 100)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels, size=11, color='white')
    ax.set_rgrids([20, 40, 60, 80], color='gray') 
    ax.set_title(radar_config['name'], size=20, weight='bold', y=1.1, color='white')
    ax.legend(loc='upper right', bbox_to_anchor=(1.5, 1.15), labelcolor='white', fontsize=12)

    return fig


# --- 5. STREAMLIT APP LAYOUT ---

st.title("⚽ St Johnstone Radar Analysis")
st.write("An interactive tool for visual player comparison using StatsBomb data.")

with st.spinner('Loading all player data from StatsBomb API... Please wait.'):
    leagues, player_df = load_all_data(USERNAME, PASSWORD)

if leagues is None or player_df is None:
    st.stop()

# Helper function for sidebar filters
def get_player_selection(player_num, full_df, league_data):
    """Creates a set of filters in the sidebar for a single player."""
    st.subheader(f"Player {player_num} Selection")
    
    league_options = sorted(league_data.keys(), key=lambda x: league_data[x]['name'])
    selected_league_id = st.selectbox(
        f'Select League (Player {player_num})', 
        options=league_options, 
        format_func=lambda x: league_data[x]['name'],
        key=f"league_{player_num}"
    )

    available_seasons = league_data[selected_league_id]['seasons']
    season_options = sorted(available_seasons.keys(), key=lambda k: available_seasons[k], reverse=True)
    selected_season_id = st.selectbox(
        f'Select Season (Player {player_num})', 
        options=season_options,
        format_func=lambda x: available_seasons[x],
        key=f"season_{player_num}"
    )
    
    filtered_df = full_df[(full_df['competition_id'] == selected_league_id) & (full_df['season_id'] == selected_season_id)].copy()
    
    if filtered_df.empty:
        st.warning("No player data for this selection.")
        return None, None, None

    team_list = sorted(filtered_df['team_name'].unique())
    team_list.insert(0, "All Teams")
    selected_team = st.selectbox(
        f'Select Team (Player {player_num})', 
        options=team_list,
        key=f"team_{player_num}"
    )
    
    if selected_team != "All Teams":
        player_pool_df = filtered_df[filtered_df['team_name'] == selected_team]
    else:
        player_pool_df = filtered_df
        
    if player_pool_df.empty:
        st.write("No players to select.")
        return None, None, None

    player_list = sorted(player_pool_df['player_name'].unique())
    default_index = 1 if player_num == 2 and len(player_list) > 1 else 0
    selected_player_name = st.selectbox(
        f'Select Player {player_num}', 
        options=player_list,
        key=f"player_{player_num}",
        index=default_index
    )
    
    if not selected_player_name:
        return None, None, None

    player_data = filtered_df[filtered_df['player_name'] == selected_player_name].iloc[0]
    return player_data, selected_league_id, selected_season_id

# Sidebar with dual filters
with st.sidebar:
    st.header("📊 Filters")
    player1_data, p1_league_id, p1_season_id = get_player_selection(1, player_df, leagues)
    st.divider()
    player2_data, p2_league_id, p2_season_id = get_player_selection(2, player_df, leagues)

# Main content area for charts
if player1_data is not None and player2_data is not None:
    player1_name = player1_data['player_name']
    player2_name = player2_data['player_name']
    
    if player1_data.equals(player2_data):
        st.error("Please select two different players or seasons for comparison.")
    else:
        st.header(f"Comparison: {player1_name} vs. {player2_name}")
        
        col_info1, col_info2 = st.columns(2)
        with col_info1:
            st.markdown(f"**{player1_name}**")
            st.markdown(f"League: **{leagues[p1_league_id]['name']}**")
            st.markdown(f"Season: **{leagues[p1_league_id]['seasons'][p1_season_id]}**")
        with col_info2:
            st.markdown(f"**{player2_name}**")
            st.markdown(f"League: **{leagues[p2_league_id]['name']}**")
            st.markdown(f"Season: **{leagues[p2_league_id]['seasons'][p2_season_id]}**")

        st.markdown("---")

        col1, col2 = st.columns(2)
        cols = [col1, col2, col1, col2, col1, col2]
        
        radar_configs = list(WINGER_RADAR_METRICS.values())
        
        for i, config in enumerate(radar_configs):
            with cols[i]:
                fig = create_comparison_radar_chart(player1_data, player2_data, config)
                st.pyplot(fig, use_container_width=True)

                doc_stream = create_word_doc_stream(player1_name, player2_name, config['name'], fig)
                st.download_button(
                    label="Download as .docx",
                    data=doc_stream,
                    file_name=f"{config['name']}_{player1_name}_vs_{player2_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{i}"
                )
else:
    st.info("Select two players from the sidebar to see the comparison.")
    
    
    
    
    
    
    
    
    
    
    
    
